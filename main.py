# coding=utf-8
"""
    @Author  : seeklife
    @Time    : 2020/12/07
    @Comment :
"""

import Tkinter as tk
from tkinter import ttk
from tkinter.filedialog import *
from tkMessageBox import askokcancel
from tkMessageBox import showerror
from docx import Document
from docx.shared import Pt   # 磅数
from docx.oxml.ns import qn  # chinese
import os
from threading import Thread
from xml.dom.minidom import parse
import xml.dom.minidom as minidom
from win32com import client as wc  # 导入doc转docx
import pythoncom                   # 解决线程无法使用Pywin32问题
# 引入处理Excel表格的方法
import rAndwExcel
import rAndwExcelB
import webbrowser

import sys

reload(sys)
sys.setdefaultencoding('utf-8')


# 找到文件夹下的所有doxc文件并获得文件名
def file_name(file_dir):
    for root, dirs, files in os.walk(file_dir):
        return files


# 遍历刚才找到的所有文件夹然后替换关键字
def change_text(old_text, new_text, document):
    all_paragraphs = document.paragraphs
    for paragraph in all_paragraphs:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

    all_tables = document.tables
    for table in all_tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # print '扫描-->'+run.text
                        if old_text in run.text:
                            # print run.text+'->'+new_text #+'-->'+chardet.detect(cell.text.encode())
                            run.text = run.text.replace(old_text, new_text)


def doc_to_docx(file1, file2):
    word = wc.Dispatch("Word.Application")  # 打开word应用程序
    doc = word.Documents.Open(file1)  # 打开word文件
    doc.SaveAs("{}x".format(file2), 12)  # 另存为后缀为".docx"的文件，其中参数12指docx文件
    doc.Close()  # 关闭原来word文件
    word.Quit()
    return "{}x".format(file)  # 返回的是一个file对象或者文件名称


def docx_to_doc(file1, file2):
    word = wc.Dispatch("Word.Application")  # 打开word应用程序
    doc = word.Documents.Open(file1)  # 打开word文件
    doc.SaveAs("{}".format(file2[:-1]), 0)  # 另存为后缀为".docx"的文件，其中参数12指docx文件
    doc.Close()  # 关闭原来word文件
    word.Quit()
    # return "{}x".format(file)  #返回的是一个file对象或者文件名称


class Application(tk.Frame):  # application对象继承tk.Frame
    def __init__(self, master=None):
        tk.Frame.__init__(self, master)  # 2.x版本不能使用super().__init__()的方式构建父类而且父类必须是新类而不是经典类
        self.master = master             # 这个master就是tkinter的窗口对象
        self.pack()
        self.window_flag = 0             # 等待窗口结束标志
        self.doc_flag = 0                # 等待转换结束标志
        self.entry_row = 1               # 输入框的行号,默认值为1docx
        self.entry_list = []             # 创建全局的entry的list
        self.entry_str = []              # 创建全局的entry的str的list
        self.entry_label = []            # 存储显示的所有label组件
        self.entry_label_A = []          # 存储显示的所有表格功能A的label组件
        self.entry_str_A = []            # 创建全局的所有表格功能A的entry的str的list
        self.entry_list_A = []           # 存储显示的所有表格功能A的entry组件
        self.entry_label_B = []          # 存储显示的所有表格功能B的label组件
        self.entry_str_B = []            # 创建全局的所有表格功能B的entry的str的list
        self.entry_list_B = []           # 存储显示的所有表格功能B的entry组件
        # self.rageValues = []           #存储抓取内容的行范围
        self.rowValues = []              # 存储抓取内容的单个行
        self.c_height = 150              # 自适应高度 docx
        self.EA_entry_row = 3            # 输入框的行号,默认值为1docx ExcelA
        self.EA_height = 150             # 自适应高度 docx
        self.EB_entry_row = 4            # 输入框的行号,默认值为ExcelB
        self.EB_height = 150             # 自适应高度 ExcelB
        self.importPath = ''
        self.exportPath = ''
        # self.xVariable = ''            #下拉菜单的值
        self.create_widgets()            # 该方法用来创建组件

    # 所有的组件都放在这个方法中进行创建
    def create_widgets(self):
        # 创建菜单栏
        self.m = tk.Menu(self)
        self.master['menu'] = self.m
        self.m_m1 = tk.Menu(self.m, tearoff=False)
        self.m_m1.add_command(label='保存', command=self.saveXml)
        # self.m_m1.add_command(label='另存为')
        self.m_m1.add_command(label='加载', command=self.loadXml)
        self.m_m1.add_command(label='导入路径', command=self.openDirectory1)
        self.m_m1.add_command(label='导出路径', command=self.openDirectory2)
        self.m_m2 = tk.Menu(self.m, tearoff=False)
        self.m_m2.add_command(label='doc转docx', command=lambda: self.createDocToDocx(1))  # 这里使用lambda是因为需要传值
        self.m_m2.add_command(label='docx转doc', command=lambda: self.createDocToDocx(4))
        self.m_m2.add_command(label='xls转xlsx', command=lambda: self.createDocToDocx(2))
        self.m_m2.add_command(label='xlsx转xls', command=lambda: self.createDocToDocx(3))
        self.m_m3 = tk.Menu(self.m, tearoff=False)
        self.m_m3.add_command(label='模式说明', command=self.goToLink)
        self.m_m4 = tk.Menu(self.m, tearoff=False)
        # 主菜单
        self.m.add_cascade(label='文件', menu=self.m_m1)
        self.m.add_cascade(label='工具', menu=self.m_m2)
        self.m.add_cascade(label='帮助', menu=self.m_m3)

        # 以下内容除一直到提交这都放在一个有滑块的可变frame中
        # 要想实现frame带滚动条必须使用canvas来实现,把frame放入canvas中给canvas添加滚动条
        # 创建画布
        # 下拉菜单
        self.fmode = tk.Frame(self)
        self.fmode.grid(row=0, column=0, columnspan=3, sticky="n" + "s" + "w" + "e", pady=5)
        tk.Label(self.fmode, text="模式选择：").pack(side='left')
        self.xVariable = StringVar()                                      # 创建变量，便于取值
        self.com = ttk.Combobox(self.fmode, textvariable=self.xVariable)  # 创建下拉菜单
        self.com.pack(side='left')                                        # 将下拉菜单绑定到窗体
        self.com["value"] = ("docx内容替换", "xls/xlsx多导一", "xls/xlsx一导多")  # 给下拉菜单设定值
        self.com["state"] = "readonly"
        self.com.current(0)                                                        # 设定下拉菜单的默认值为第1个
        self.com.bind("<<ComboboxSelected>>", self.modeCombbox)
        self.combboxStr = StringVar()
        self.lcombox = tk.Label(self.fmode, textvariable=self.combboxStr, fg='blue')
        self.combboxStr.set('当前模式：' + self.com.get())
        self.lcombox.pack(side='right')                                            # 设置滑块的位置

        # --------------------------------需要改变的内容--------------------------------------------1
        self.canvas = Canvas(self, width=420, height=140,highlightthickness=0)     # 420 scrollregion=(0,0,self.c_height,self.c_height)默认就是超过长或者宽滑块自动显示
        self.canvas.grid(row=1, column=0, columnspan=3, ipady=2, pady=5, sticky=tk.N)
        self.f1 = tk.Frame(self.canvas, highlightthickness=1, highlightbackground='black', highlightcolor='black')
        self.f1.place(width=400, height=145)  # 这里才能设置frame在画布中的大小,有了create_window这里不会起作用了
        # 创建竖直滚动条
        sb1 = Scrollbar(self.canvas, orient=VERTICAL)
        sb1.place(x=400, width=20, height=145)
        sb1.configure(command=self.canvas.yview)
        self.canvas.config(yscrollcommand=sb1.set)  # 设置滚动条,xscrollcommand=sb2.set
        self.canvas.create_window((200, 0), window=self.f1, anchor=N)  # 表示把组件放入到canvas中成为它的一部分,这样滑动滚动条也可以滑动frame

        # 第一行
        tk.Label(self.f1, text="-----old_text------").grid(row=0, column=0, ipadx=30, sticky=tk.E, ipady=2, pady=5)
        tk.Label(self.f1, text="：").grid(row=0, column=1, ipadx=5, ipady=2, sticky=tk.E, pady=5)
        tk.Label(self.f1, text="-------new_text------").grid(row=0, column=2, ipadx=25, ipady=2, sticky=tk.E)

        # 第一个输入框,添加全局列表中
        self.sl1 = StringVar()
        self.sl1 = '1.'
        l1 = tk.Label(self.f1, text=self.sl1)
        self.entry_label.append(l1)
        l1.grid(row=1, column=0, ipadx=5, ipady=2, sticky=tk.W, pady=5)
        self.s1 = StringVar()
        self.entry_str.append(self.s1)
        self.e1 = tk.Entry(self.f1, textvariable=self.s1)
        self.entry_list.append(self.e1)
        self.e1.grid(row=1, column=0, ipadx=1, sticky=tk.E, ipady=2, pady=5)
        l2 = tk.Label(self.f1, text="：")
        self.entry_label.append(l2)
        l2.grid(row=1, column=1, ipadx=5, ipady=2, sticky=tk.E, pady=5)
        self.s2 = StringVar()
        self.entry_str.append(self.s2)
        self.e2 = tk.Entry(self.f1, textvariable=self.s2)
        self.entry_list.append(self.e2)
        self.e2.grid(row=1, column=2, ipadx=25, ipady=2, sticky=tk.E, padx=5)
        # ----------------------改变内容结束--------------------------------------------------------1

        # --------------------------------需要改变的内容--------------------------------------------2
        self.canvas1 = Canvas(self, width=420, height=140,highlightthickness=0)                   # scrollregion=(0,0,self.c_height,self.c_height)默认就是超过长或者宽滑块自动显示
        self.canvas1.grid(row=1, column=0, columnspan=3, ipady=2, pady=5, sticky=tk.N)
        self.canvas1.grid_forget()
        self.f2 = tk.Frame(self.canvas1, highlightthickness=1, highlightbackground='black', highlightcolor='black')
        # canvas.create_window((200,100),window=self.f1)
        self.f2.place(width=400, height=145)  # 这里才能设置frame在画布中的大小,有了create_window这里不会起作用了
        # 创建竖直滚动条
        sb2 = Scrollbar(self.canvas1, orient=VERTICAL)
        sb2.place(x=400, width=20, height=145)
        sb2.configure(command=self.canvas1.yview)
        self.canvas1.config(yscrollcommand=sb2.set)  # 设置滚动条,xscrollcommand=sb2.set
        self.canvas1.create_window((200, 0), window=self.f2, anchor=N)  # 表示把组件放入到canvas中成为它的一部分,这样滑动滚动条也可以滑动frame

        # 第一行
        self.sl2_1 = StringVar()
        self.sl2_1 = '复制到开始行数：'
        l2_1 = tk.Label(self.f2, text=self.sl2_1)
        # self.entry_label_A.append(l2_1)
        l2_1.grid(row=0, column=0, columnspan=2, ipadx=5, ipady=2, sticky=tk.W, pady=5)
        self.s2_1 = StringVar()
        # self.entry_str_A.append(self.s2_1)
        self.e2_1 = tk.Entry(self.f2, textvariable=self.s2_1, width=5)
        # self.entry_list_A.append(self.e2_1)
        self.e2_1.grid(row=0, column=2, ipadx=5, sticky=tk.W, ipady=2, pady=3)
        l2_2 = tk.Label(self.f2, text="复制到开始列数：")
        # self.entry_label_A.append(l2_2)
        l2_2.grid(row=1, column=0, columnspan=2, ipadx=5, ipady=2, pady=3, sticky=tk.W)
        self.s2_2 = StringVar()
        # self.entry_str_A.append(self.s2_2)
        self.e2_2 = tk.Entry(self.f2, textvariable=self.s2_2, width=5)
        # self.entry_list_A.append(self.e2_2)
        self.e2_2.grid(row=1, column=2, ipadx=5, sticky=tk.W, ipady=2, pady=3)

        tk.Label(self.f2, text="-----内容名称-----").grid(row=2, column=0, columnspan=3, ipadx=5, sticky='e' + 'w', ipady=2,pady=5)
        tk.Label(self.f2, text="-行-").grid(row=2, column=3, ipadx=5, ipady=2, sticky='e' + 'w', pady=5)
        tk.Label(self.f2, text="-列-").grid(row=2, column=5, ipadx=25, ipady=2, sticky='e' + 'w')

        # 选择坐标
        self.sl2_3 = StringVar()
        self.sl2_3 = '1.'
        l2_3 = tk.Label(self.f2, text=self.sl2_3)
        self.entry_label_A.append(l2_3)
        l2_3.grid(row=3, column=0, ipadx=5, ipady=2, sticky=tk.W, pady=5)
        # 坐标内容提示
        self.s2_3 = StringVar()
        self.entry_str_A.append(self.s2_3)
        self.e2_3 = tk.Entry(self.f2, textvariable=self.s2_3, width=15)
        self.entry_list_A.append(self.e2_3)
        self.e2_3.grid(row=3, column=1, ipadx=5, sticky=tk.E, ipady=2, pady=5)

        l2_4 = tk.Label(self.f2, text="：")
        self.entry_label_A.append(l2_4)
        l2_4.grid(row=3, column=2, ipadx=5, ipady=2, sticky='e' + 'w', pady=5)
        # 爬取内容坐标的行
        self.s2_4 = StringVar()
        self.entry_str_A.append(self.s2_4)
        self.e2_4 = tk.Entry(self.f2, textvariable=self.s2_4, width=5)
        self.entry_list_A.append(self.e2_4)
        self.e2_4.grid(row=3, column=3, ipadx=5, ipady=2, sticky=tk.E)

        l2_5 = tk.Label(self.f2, text="-")
        self.entry_label_A.append(l2_5)
        l2_5.grid(row=3, column=4, ipadx=5, ipady=2, pady=5)
        # 爬取内容的列数
        self.s2_5 = StringVar()
        self.entry_str_A.append(self.s2_5)
        self.e2_5 = tk.Entry(self.f2, textvariable=self.s2_5, width=5)
        self.entry_list_A.append(self.e2_5)
        self.e2_5.grid(row=3, column=5, ipadx=5, ipady=2)
        # ----------------------改变内容结束----------------------------------------------2

        # --------------------------------需要改变的内容-----------------------------------3
        self.canvas2 = Canvas(self, width=420, height=140,highlightthickness=0)  # ,scrollregion=(0,0,self.c_height,self.c_height)默认就是超过长或者宽滑块自动显示
        self.canvas2.grid(row=1, column=0, columnspan=3, ipady=2, pady=5, sticky=tk.N)
        self.canvas2.grid_forget()
        self.f3 = tk.Frame(self.canvas2, highlightthickness=1, highlightbackground='black', highlightcolor='black')
        self.f3.place(width=400, height=145)  # 这里才能设置frame在画布中的大小,有了create_window这里不会起作用了
        # 创建竖直滚动条
        sb3 = Scrollbar(self.canvas2, orient=VERTICAL)
        sb3.place(x=400, width=20, height=145)
        sb3.configure(command=self.canvas2.yview)
        self.canvas2.config(yscrollcommand=sb3.set)  # 设置滚动条,xscrollcommand=sb3.set
        self.canvas2.create_window((200, 0), window=self.f3, anchor=N)  # 表示把组件放入到canvas中成为它的一部分,这样滑动滚动条也可以滑动frame

        # 第一行
        self.strB_1 = StringVar()
        self.strB_1 = '要抓取内容的行：'
        lB_1 = tk.Label(self.f3, text=self.strB_1)
        lB_1.grid(row=0, column=0, columnspan=2, ipadx=5, ipady=2, sticky=tk.W, pady=5)
        self.strB_2 = StringVar()
        self.eB_1 = tk.Entry(self.f3, textvariable=self.strB_2, width=30)
        # self.entry_list_A.append(self.e2_1)
        self.eB_1.grid(row=0, column=2, columnspan=3, ipadx=5, sticky=tk.W, ipady=2, pady=5, padx=5)
        lB_2 = tk.Label(self.f3, text="文件重命名：")
        # self.entry_label_A.append(l2_2)
        lB_2.grid(row=1, column=0, columnspan=2, ipadx=5, ipady=2, pady=5, sticky=tk.W)
        self.strB_3 = StringVar()
        # self.entry_str_A.append(self.s2_2)
        self.eB_2 = tk.Entry(self.f3, textvariable=self.strB_3, width=5)
        # self.entry_list_A.append(self.e2_2)
        self.eB_2.grid(row=1, column=2, columnspan=3, ipadx=5, sticky=tk.W, ipady=2, pady=5, padx=5)

        tk.Label(self.f3, text="列名").grid(row=2, column=0, columnspan=2, ipadx=1, sticky='e' + 'w', ipady=2, pady=5)
        tk.Label(self.f3, text="列数").grid(row=2, column=2, ipadx=1, ipady=2, sticky='e' + 'w', pady=5)
        tk.Label(self.f3, text="复制到行").grid(row=2, column=3, ipadx=5, ipady=2, sticky='e' + 'w')
        tk.Label(self.f3, text="复制到列").grid(row=2, column=4, ipadx=5, ipady=2, sticky='e' + 'w')

        # 需要添加的内容
        # 选择坐标
        self.strB_4 = StringVar()
        self.strB_4 = '1.'
        lB_2 = tk.Label(self.f3, text=self.strB_4)
        self.entry_label_B.append(lB_2)
        lB_2.grid(row=3, column=0, ipadx=1, ipady=2, sticky=tk.W, pady=2)
        # 列名的列
        self.strB_cName = StringVar()
        self.entry_str_B.append(self.strB_cName)
        self.eB_3 = tk.Entry(self.f3, textvariable=self.strB_cName, width=15)
        self.entry_list_B.append(self.eB_3)
        self.eB_3.grid(row=3, column=1, ipadx=1, sticky='e' + 'w', ipady=2, pady=2, padx=2)

        # 爬取的列坐标
        self.strB_5 = StringVar()
        self.entry_str_B.append(self.strB_5)
        self.eB_4 = tk.Entry(self.f3, textvariable=self.strB_5, width=5)
        self.entry_list_B.append(self.eB_4)
        self.eB_4.grid(row=3, column=2, ipadx=1, ipady=2, sticky='e' + 'w', padx=2)

        # 复制到的行
        self.strB_6 = StringVar()
        self.entry_str_B.append(self.strB_6)
        self.eB_5 = tk.Entry(self.f3, textvariable=self.strB_6, width=5)
        self.entry_list_B.append(self.eB_5)
        self.eB_5.grid(row=3, column=3, ipadx=1, ipady=2, sticky='e' + 'w', padx=2)

        # 复制到的列
        self.strB_7 = StringVar()
        self.entry_str_B.append(self.strB_7)
        self.eB_6 = tk.Entry(self.f3, textvariable=self.strB_7, width=5)
        self.entry_list_B.append(self.eB_6)
        self.eB_6.grid(row=3, column=4, ipadx=1, ipady=2, sticky='e' + 'w', padx=4)
        # ----------------------改变内容结束----------------------------------------------3

        tk.Button(self, text='添加', command=lambda: self.b_add(0), height=1, width=5).grid(row=2, column=0, columnspan=4, pady=5)
        # 第20行
        tk.Button(self, text='提交', command=self.b1_run, height=1, width=5).grid(row=3, column=0, columnspan=4, pady=5)
        # 第21行
        tk.Label(self, text=u'当前导入路径：').grid(row=4, column=0, ipadx=5, ipady=2, sticky=tk.W)
        self.pathEntry1_s = StringVar()
        self.pathEntry1 = tk.Entry(self, textvariable=self.pathEntry1_s, width=70)
        self.pathEntry1.grid(row=4, column=1, columnspan=2, ipadx=5, ipady=2, sticky=tk.W)
        # 第22行
        tk.Label(self, text=u'当前导出路径：').grid(row=5, column=0, ipadx=5, ipady=2, sticky=tk.W)
        self.pathEntry2_s = StringVar()
        self.pathEntry2 = tk.Entry(self, textvariable=self.pathEntry2_s, width=70)
        self.pathEntry2.grid(row=5, column=1, columnspan=2, ipadx=5, ipady=2, sticky=tk.W)
        # 窗口关闭事件
        # 点击右上关闭按钮退出
        self.master.protocol('WM_DELETE_WINDOW', self.on_closing)

    # 下拉菜单点击事件
    # 每一次的下拉菜单点击事件路径必须清空，路径指导入导出路径
    def modeCombbox(self, event):
        print '下拉菜单被点击了'
        print self.com.get()
        if self.com.get() == 'docx内容替换':
            print '0'
            self.combboxStr.set('当前模式：' + self.com.get())
            self.canvas.grid(row=1, column=0, columnspan=3, ipady=2, pady=5, sticky=tk.N)
            self.canvas1.grid_forget()
            self.canvas2.grid_forget()
            self.pathEntry1_s.set('')
            self.pathEntry2_s.set('')
            self.importPath = ''
            self.exportPath = ''
        elif self.com.get() == 'xls/xlsx内容导入A':
            print '1'
            self.combboxStr.set('当前模式：' + self.com.get())
            self.canvas.grid_forget()
            self.canvas2.grid_forget()
            self.canvas1.grid(row=1, column=0, columnspan=3, ipady=2, pady=5, sticky=tk.N)
            self.pathEntry1_s.set('')
            self.pathEntry2_s.set('')
            self.importPath = ''
            self.exportPath = ''
        elif self.com.get() == 'xls/xlsx内容导入B':
            print '2'
            self.combboxStr.set('当前模式：' + self.com.get())
            self.canvas.grid_forget()
            self.canvas1.grid_forget()
            self.canvas2.grid(row=1, column=0, columnspan=3, ipady=2, pady=5, sticky=tk.N)
            self.pathEntry1_s.set('')
            self.pathEntry2_s.set('')
            self.importPath = ''
            self.exportPath = ''

    # 所有该应用的其他方法都在此处创建
    # 关闭窗口询问是否退出
    def on_closing(self):
        # 这里应该提醒用户是否保存当前配置
        if askokcancel('退出', u"你确定要退出吗?"):
            self.master.destroy()

    # 处理提示完成
    def complate_info(self):
        askokcancel('提示', u"处理完成！")

    # 消息错误提示框
    def message_error(self):
        showerror("错误", u"你的导入或者导出路径不存在")

    # 行和列不能为空
    def message_error_rc(self):
        showerror("错误", u"行和列不能为空")

    # 行和列不能为空
    def message_error_B(self):
        showerror("错误", u"行和重命名不能为空")

    # 添加按钮点击事件处理函数
    def b_add(self, s_length):
        # 点击按钮时查看当前的模式
        if self.com.get() == 'docx内容替换':
            print 'Docx添加'  # 每次添加必须知道是第几行
            self.entry_row += 1  # entry行号
            if s_length != 0:
                print '加载添加' + str(s_length)
                self.c_height += s_length  # 获取frame增加高度
                # 设置滚动的滚动范围，发现范围显示不全((0,0,self.c_height,self.c_height))或者canvas.bbox('all')
                self.canvas.configure(scrollregion=((0, 0, self.c_height, self.c_height)))  # 将滚动区域设置为画布边界
            else:
                print '按钮添加' + str(s_length)
                self.c_height += (self.f1.winfo_height() + 30 - self.c_height)  # 获取frame增加高度
                # 设置滚动的滚动范围，发现范围显示不全((0,0,self.c_height,self.c_height))或者canvas.bbox('all')
                self.canvas.configure(scrollregion=((0, 0, self.c_height, self.c_height)))  # 将滚动区域设置为画布边界
            # 增加新的组件内容，同时输入框的引用增加，输入框字符串的引用增加
            self.sl1 = StringVar()
            self.sl1 = str(self.entry_row) + '.'
            l1 = tk.Label(self.f1, text=self.sl1)
            self.entry_label.append(l1)
            l1.grid(row=self.entry_row, column=0, ipadx=5, ipady=2, sticky=tk.W, pady=5) #--->
            self.s1 = StringVar()
            self.entry_str.append(self.s1)
            self.e1 = tk.Entry(self.f1, textvariable=self.s1)
            self.entry_list.append(self.e1)
            self.e1.grid(row=self.entry_row, column=0, ipadx=1, sticky=tk.E, ipady=2, pady=5)
            l2 = tk.Label(self.f1, text="：")
            self.entry_label.append(l2)
            l2.grid(row=self.entry_row, column=1, ipadx=5, ipady=2, sticky=tk.E, pady=5)
            self.s2 = StringVar()
            self.entry_str.append(self.s2)
            self.e2 = tk.Entry(self.f1, textvariable=self.s2)
            self.entry_list.append(self.e2)
            self.e2.grid(row=self.entry_row, column=2, ipadx=25, ipady=2, sticky=tk.E, padx=5)
            # 改变之修改相对位置,再调用一次create_window
            # self.canvas.create_window((200,(75+(self.f1.winfo_height() + 96 - self.c_height))),window=self.f1)
        elif self.com.get() == 'xls/xlsx内容导入A':
            print 'EA添加'  # 每次添加必须知道是第几行
            self.EA_entry_row += 1  # entry行号
            # print '每次增加距离-->' + str(self.f2.winfo_height())
            # 判断一下是否是加载添加操作
            if s_length != 0:
                print '加载添加' + str(s_length)
                self.EA_height += s_length  # 获取frame增加高度
                # 设置滚动的滚动范围，发现范围显示不全((0,0,self.c_height,self.c_height))或者canvas.bbox('all')
                self.canvas1.configure(scrollregion=((0, 0, self.EA_height, self.EA_height)))  # 将滚动区域设置为画布边界
            else:
                print '按钮添加' + str(s_length)
                self.EA_height += (self.f2.winfo_height() + 30 - self.EA_height)  # 获取frame增加高度
                # 设置滚动的滚动范围，发现范围显示不全((0,0,self.c_height,self.c_height))或者canvas.bbox('all')
                self.canvas1.configure(scrollregion=((0, 0, self.EA_height, self.EA_height)))  # 将滚动区域设置为画布边界
            # 增加新的组件内容，同时输入框的引用增加，输入框字符串的引用增加
            # 选择坐标
            self.sl2_3 = StringVar()
            self.sl2_3 = str(self.EA_entry_row - 2) + '.'
            l2_3 = tk.Label(self.f2, text=self.sl2_3)
            self.entry_label_A.append(l2_3)
            l2_3.grid(row=self.EA_entry_row, column=0, ipadx=1, ipady=2, sticky=tk.W, pady=5)
            self.s2_3 = StringVar()
            self.entry_str_A.append(self.s2_3)
            self.e2_3 = tk.Entry(self.f2, textvariable=self.s2_3, width=15)
            self.entry_list_A.append(self.e2_3)
            self.e2_3.grid(row=self.EA_entry_row, column=1, ipadx=5, sticky=tk.E, ipady=2, pady=5)
            l2_4 = tk.Label(self.f2, text="：")
            self.entry_label_A.append(l2_4)
            l2_4.grid(row=self.EA_entry_row, column=2, ipadx=5, ipady=2, sticky='e' + 'w', pady=5)
            self.s2_4 = StringVar()
            self.entry_str_A.append(self.s2_4)
            self.e2_4 = tk.Entry(self.f2, textvariable=self.s2_4, width=5)
            self.entry_list_A.append(self.e2_4)
            self.e2_4.grid(row=self.EA_entry_row, column=3, ipadx=5, ipady=2, sticky=tk.E)
            l2_5 = tk.Label(self.f2, text="-")
            self.entry_label_A.append(l2_5)
            l2_5.grid(row=self.EA_entry_row, column=4, ipadx=5, ipady=2, pady=5)
            self.s2_5 = StringVar()
            self.entry_str_A.append(self.s2_5)
            self.e2_5 = tk.Entry(self.f2, textvariable=self.s2_5, width=5)
            self.entry_list_A.append(self.e2_5)
            self.e2_5.grid(row=self.EA_entry_row, column=5, ipadx=5, ipady=2)
        elif self.com.get() == 'xls/xlsx内容导入B':
            # print '3'
            self.EB_entry_row += 1  # entry行号
            # print '每次增加距离-->' + str(self.f2.winfo_height())
            # 判断一下是否是加载添加操作
            if s_length != 0:
                print '加载添加' + str(s_length)
                self.EB_height += s_length  # 获取frame增加高度
                # 设置滚动的滚动范围，发现范围显示不全((0,0,self.c_height,self.c_height))或者canvas.bbox('all')
                self.canvas2.configure(scrollregion=((0, 0, self.EB_height, self.EB_height)))  # 将滚动区域设置为画布边界
            else:
                print '按钮添加' + str(s_length)
                self.EB_height += (self.f3.winfo_height() + 30 - self.EB_height)  # 获取frame增加高度
                # 设置滚动的滚动范围，发现范围显示不全((0,0,self.c_height,self.c_height))或者canvas.bbox('all')
                self.canvas2.configure(scrollregion=((0, 0, self.EB_height, self.EB_height)))  # 将滚动区域设置为画布边界
            # 增加新的组件内容，同时输入框的引用增加，输入框字符串的引用增加
            # 选择坐标
            self.strB_4 = StringVar()
            self.strB_4 = str(self.EB_entry_row - 3) + '.'
            lB_2 = tk.Label(self.f3, text=self.strB_4)
            self.entry_label_B.append(lB_2)
            lB_2.grid(row=self.EB_entry_row, column=0, ipadx=1, ipady=2, sticky=tk.W, pady=2)
            # 列名的列
            self.strB_cName = StringVar()
            self.entry_str_B.append(self.strB_cName)
            self.eB_3 = tk.Entry(self.f3, textvariable=self.strB_cName, width=15)
            self.entry_list_B.append(self.eB_3)
            self.eB_3.grid(row=self.EB_entry_row, column=1, ipadx=1, sticky='e' + 'w', ipady=2, pady=2, padx=2)

            # 爬取的列坐标
            self.strB_5 = StringVar()
            self.entry_str_B.append(self.strB_5)
            self.eB_4 = tk.Entry(self.f3, textvariable=self.strB_5, width=5)
            self.entry_list_B.append(self.eB_4)
            self.eB_4.grid(row=self.EB_entry_row, column=2, ipadx=1, ipady=2, sticky='e' + 'w', padx=2)

            # 复制到的行
            self.strB_6 = StringVar()
            self.entry_str_B.append(self.strB_6)
            self.eB_5 = tk.Entry(self.f3, textvariable=self.strB_6, width=5)
            self.entry_list_B.append(self.eB_5)
            self.eB_5.grid(row=self.EB_entry_row, column=3, ipadx=1, ipady=2, sticky='e' + 'w', padx=2)

            # 复制到的列
            self.strB_7 = StringVar()
            self.entry_str_B.append(self.strB_7)
            self.eB_6 = tk.Entry(self.f3, textvariable=self.strB_7, width=5)
            self.entry_list_B.append(self.eB_6)
            self.eB_6.grid(row=self.EB_entry_row, column=4, ipadx=1, ipady=2, sticky='e' + 'w', padx=4)

    # docx提交按钮点击事件处理函数
    def b1_run(self):
        # 创建新线程来执行任务
        # 根据模式来选择不同任务
        if self.com.get() == 'docx内容替换':
            # 在按下按钮之前要判断导入导出路径是否填写，是否有效
            if (self.exportPath != '' and self.exportPath != None) and (
                    self.importPath != '' and self.importPath != None) and (
                    os.path.exists(self.importPath) and os.path.exists(self.exportPath)):
                self.window_flag = 0  # 表示等待弹窗保持开启状态
                # 用主线程创建提示窗口
                self.createWait()
                print "开始执行doc替换操作"
                self.thread_task(self.deal_task)
            else:
                self.message_error()
        elif self.com.get() == 'xls/xlsx内容导入A':
            # 在按下按钮之前要判断导入导出路径是否填写，是否有效
            # 除了处理任务不一样其他代码均可复用
            if (self.exportPath != '' and self.exportPath != None) and (
                    self.importPath != '' and self.importPath != None) and (
                    os.path.exists(self.importPath) and os.path.exists(self.exportPath)):
                # 执行之前还需要额外判断一下行和列是否为空
                if (self.s2_1.get() != '' and self.s2_2.get() != '') and (
                        self.s2_1.get() != None and self.s2_2.get() != None):
                    self.window_flag = 0  # 表示等待弹窗保持开启状态
                    # 用主线程创建提示窗口
                    self.createWait()
                    self.thread_task(self.deal_Excel_A)
                else:
                    self.message_error_rc()
            else:
                self.message_error()
        # 模式B还未开发完
        elif self.com.get() == 'xls/xlsx内容导入B':
            # 在按下按钮之前要判断导入导出路径是否填写，是否有效
            # 除了处理任务不一样其他代码均可复用
            if (self.exportPath != '' and self.exportPath != None) and (
                    self.importPath != '' and self.importPath != None) and (
                    os.path.exists(self.importPath) and os.path.exists(self.exportPath)):
                # 执行之前还需要额外判断一下抓取的行和文件重命名是否为空否则警告
                if (self.strB_2.get() != '' and self.strB_3.get() != '') and (
                        self.strB_2.get() != None and self.strB_3.get() != None):
                    self.window_flag = 0  # 表示等待弹窗保持开启状态
                    # 用主线程创建提示窗口
                    self.createWait()
                    self.thread_task(self.deal_Excel_B)
                else:
                    self.message_error_rc()
            else:
                self.message_error()
        # 这里调用一下延时递归,每200ms之后触发,检测标志位一旦符合条件结束递归即return
        self.taskOver(self.master)

    # 动态根据替换数量执行任务
    # 在执行之前判断一下路径是否存在
    # 判断路径存在后要对文件夹内部文件进行筛选，只选择docx文件，或者进行转换
    def deal_task(self):
        # print u"新线程开始执行任务"
        words = file_name(self.importPath)
        for words_name in words:
            print words_name
            # 只有后缀为docx的文件可以继续执行
            if words_name.find('.docx') != -1:
                # PackageNotFoundError,捕获一下该异常
                try:
                    document = Document(self.importPath + '/' + words_name)
                    document.styles['Normal'].font.name = u"仿宋_GB2312"
                    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u"仿宋_GB2312")
                    document.styles['Normal'].font.size = Pt(12)
                    # 进行替换操作
                    # old_text和new_text都是动态传入
                    # 遍历entry_list和entry_str来执行动态操作，偶是old奇数是new
                    for num in range(0, len(self.entry_list) - 1, 2):
                        # 输入框任一一栏没有输入内容不会执行替换操作
                        if self.entry_list[num].get() != '' and self.entry_list[num + 1].get() != '':
                            change_text(self.entry_list[num].get().strip(), self.entry_list[num + 1].get().strip(),
                                        document)
                            print 'old_text: ' + self.entry_list[num].get()
                            print 'new_next: ' + self.entry_list[num + 1].get()
                    document.save(self.exportPath + "/" + "auto" + words_name)
                except Exception as e:
                    print '打开docx文件失败'
        # 关闭等待弹窗,主线程不能使用其他线程创建的对象属性了，要关闭窗口只能销毁线程
        self.window_flag = 1
        # 只要在多线程涉及到和窗口有关的api一律会出现卡死情况

    # 爬取表格内容并复制到表格力功能A
    # 多个表格内容复制到另一个表，一个表的内容就是一行
    def deal_Excel_A(self):
        print u"新线程开始执行任务"
        print '用户输入行数' + self.s2_1.get()
        row = int(self.s2_1.get()) + 1  # 获得用户输入的开始行数
        print '用户输入列数' + self.s2_2.get()
        col = int(self.s2_2.get()) + 1  # 获得用户输入的开始列数
        list_col = []  # 一个人的表格内容
        list_entry = self.entry_list_A  # 获得用户输入的行和列
        list_entry_str = self.entry_str_A  # 获得用户输入的行和列
        allFilesNames = rAndwExcel.read_file_name(self.importPath)
        i = 1
        for fileName in allFilesNames:
            # print fileName
            # fileName = fileName.decode('gbk').encode('utf-8')
            print fileName
            list_col = rAndwExcel.rwExecel(self.importPath + '/' + fileName
                                           , list_col, list_entry)
            i += 1
            print '遍历的多表数目' + str(i)
        # 读内容到列表完毕后进行统一写入操作
        rAndwExcel.wExecel(self.exportPath, row, len(list_entry) / 3, col, list_col)
        # 关闭等待弹窗,主线程不能使用其他线程创建的对象属性了，要关闭窗口只能销毁线程
        self.window_flag = 1
        # 只要在多线程涉及到和窗口有关的api一律会出现卡死情况

    # 模式B的提交按钮处理函数
    def deal_Excel_B(self):
        # 初始化一下结束flag,0表示结束,1表示未结束
        self.window_flag = 0
        self.rowValues = []
        print u"新线程开始执行B模式任务"
        print '抓取内容的行' + self.strB_2.get()
        # 首先对抓去内容进行逗号分割,得到一个列表
        getInputValues = self.strB_2.get().split(',')
        # 对此数组进行遍历，查看是否有分隔符-存在
        print getInputValues
        for value in getInputValues:
            print '行的范围' + value
            if value.find('-') != -1:
                # 继续进行分隔把数据放入列表
                rages = value.split('-')
                # self.rageValues.append(rage)
                j = int(int(rages[1]) - int(rages[0])) + 1
                # i = rages[0]
                print '范围的行数' + str(j)
                for num in range(0, j):
                    num += int(rages[0])  # 从3开始 7结束
                    # print '抓取内容范围' + str(num)             #抓取内容的范围函数
                    print '需要抓取的行：' + str(num)
                    self.rowValues.append(num)  # 添加到行列表中
            else:
                self.rowValues.append(value)
        # 获取以哪一行进行重命名
        self.reNameCol = int(self.strB_3.get())
        # 获取要抓取内容的列数，获取抓取列数的名字，复制的行和列0,1,2,3,
        list_entry = self.entry_list_B
        # 保存行的数据列表
        self.rowValues
        # 抓取的列  列名 列数 复制道坐标（行，列）
        list_entry

        # 由于xlwt无法成功复制格式，使用openpyxl方式进行xls文件的操作
        rAndwExcelB.rwExecel_B(self.importPath
                               , self.exportPath, self.rowValues, self.reNameCol, list_entry)
        # 关闭等待弹窗,主线程不能使用其他线程创建的对象属性了，要关闭窗口只能销毁线程
        self.window_flag = 1
        # 只要在多线程涉及到和窗口有关的api一律会出现卡死情况

    # 开始执行转换工作，此时开辟了新线程故不能进行弹窗
    def docTodocxStart(self, mode):
        # 初始化一下结束flag,1表示结束,0表示未结束
        self.doc_flag = 0
        # 在开始转换工作开始时检测路径的正确性，如果不正确弹出窗口警告
        if (self.sdocx1.get() != '' and self.sdocx1.get() != None) and (
                self.sdocx2.get() != '' and self.sdocx2.get() != None):
            # 首先保证不为空不为空字符
            flag1 = os.path.exists(self.sdocx1.get())
            flag2 = os.path.exists(self.sdocx2.get())
            # 判断路径存在，还需要判断是否是doc文件
            if flag1 and flag2:
                # print '文件路径存在'
                pythoncom.CoInitialize()
                a = 1
                fileNames = file_name(self.sdocx1.get())  # 获得所有文件名
                filePath = self.sdocx2.get()
                for fileName in fileNames:  # 遍历文件夹下的所有文件
                    # 先判断是否有xls然后再判断xlsx，只要不是xls直接跳过
                    # fileName = fileName.decode('gb2312').encode('utf-8')
                    print '可能要处理的文件名称' + fileName + '模式' + str(mode)

                    # 首先判断模式,1 doc -> docx 2 xls -> xlsx 3 xlsx -> xls
                    if mode == 1:
                        if fileName.find('doc') != -1:
                            # 只处理xls文件,xlsx跳过不处理
                            if fileName.find('docx') != -1:
                                continue
                            print '进入doc处理循环'
                            file1 = self.sdocx1.get() + "/" + fileName
                            print 'file1:' + file1
                            file1 = file1.decode('utf-8').encode('gb2312')
                            # print 'gb2312编码的file1: ' + file1
                            file2 = filePath + "/" + fileName
                            print 'file2:' + file2
                            file2 = file2.decode('utf-8').encode('gb2312')
                            doc_to_docx(file1, file2)
                            a += 1
                            print a
                    elif mode == 2:
                        if fileName.find('xls') != -1:
                            # 只处理xls文件,xlsx跳过不处理
                            if fileName.find('xlsx') != -1:
                                continue
                            file1 = self.sdocx1.get() + "/" + fileName
                            print 'file1:' + file1
                            # file1 = file1.decode('utf-8').encode('gb2312')
                            # print 'gb2312编码的file1: ' + file1
                            file2 = filePath + "/" + fileName
                            print 'file2:' + file2
                            # file2 = file2.decode('utf-8').encode('gb2312')
                            # 调用方法完成转换
                            rAndwExcelB.xlstoxlsx(file1, file2)
                    elif mode == 3:
                        if fileName.find('xlsx') != -1:
                            # 只处理xls文件,xlsx跳过不处理
                            file1 = self.sdocx1.get() + "/" + fileName
                            print 'file1:' + file1
                            # file1 = file1.decode('utf-8').encode('gb2312')
                            # print 'gb2312编码的file1: ' + file1
                            file2 = filePath + "/" + fileName
                            print 'file2:' + file2
                            # file2 = file2.decode('utf-8').encode('gb2312')
                            # 调用方法完成转换
                            rAndwExcelB.xlsxtoxls(file1, file2)
                    elif mode == 4:
                        if fileName.find('docx') != -1:
                            print '进入docx处理循环'
                            file1 = self.sdocx1.get() + "/" + fileName
                            print 'file1:' + file1
                            file1 = file1.decode('utf-8').encode('gb2312')
                            # print 'gb2312编码的file1: ' + file1
                            file2 = filePath + "/" + fileName
                            print 'file2:' + file2
                            file2 = file2.decode('utf-8').encode('gb2312')
                            docx_to_doc(file1, file2)
                # print "完成！"
                # 完成之后进行弹窗显示,在多线程执行弹窗就会卡死,所以换成定时函数来执行弹窗操作
                self.doc_flag = 1  # 此时表示任务结束
            else:
                # print '路径有错误'
                # print 'flag1:'+flag1+'  '+'flag2:'+flag2
                self.doc_flag = 2
        else:
            # print '路径有错误'
            # print 'flag1:'+flag1+'  '+'flag2:'+flag2
            self.doc_flag = 2

    # doc转docx的提交按钮事件处理函数
    def docTodocxTask(self, mode):
        # 启动进度条,在提交按钮下显示进度条
        self.pb.grid(row=3, column=0, columnspan=3, ipadx=5, ipady=2, pady=5, padx=5)
        self.startProgressBar()  # 启动进度条
        # 此时窗口变为不可选中状态且永远置顶
        self.docxWindow.attributes("-disabled", True)
        self.docxWindow.attributes("-topmost", True)
        # 开启线程处理dox转docx
        self.thread_task(lambda: self.docTodocxStart(mode))
        # 开启检测结束定时函数
        self.docTodocxOver(self.master)
        # print '提交按钮事件开始处理'

    # 弹出doc转docx的窗口
    def createDocToDocx(self, mode):
        # print u"创建窗口方法执行"
        self.docxWindow = tk.Toplevel(self.master)
        # print u"创建了窗口对象"
        width = 350
        height = 150
        # 获取屏幕尺寸以计算布局参数，使窗口居屏幕中央
        screenwidth = self.docxWindow.winfo_screenwidth()
        screenheight = self.docxWindow.winfo_screenheight()
        alignstr = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
        # 根据模式进行标题选择
        if mode == 1:
            self.docxWindow.title('doc转docx工具')
        elif mode == 2:
            self.docxWindow.title('xls转xlsx工具')
        elif mode == 3:
            self.docxWindow.title('xlsx转xls工具')
        elif mode == 4:
            self.docxWindow.title('docx转doc工具')
        self.docxWindow.geometry(alignstr)
        self.docxWindow.resizable(0, 0)  # 防止用户调尺寸
        tk.Label(self.docxWindow, text='导入：').grid(row=0, column=0, ipadx=5, ipady=2, pady=5)
        self.sdocx1 = StringVar()  # 设置所选路径进行显示
        self.edocx1 = tk.Entry(self.docxWindow, textvariable=self.sdocx1, width=30).grid(row=0, column=1, ipadx=5,ipady=2,pady=5)  # 设置entry为不可选不可写
        tk.Button(self.docxWindow, text='选择', command=lambda: self.openDirectory3(1)).grid(row=0, column=2, ipadx=5,ipady=2, pady=5, padx=5,sticky=tk.E)
        tk.Label(self.docxWindow, text='导出：').grid(row=1, column=0, ipadx=5, ipady=2, pady=5)
        self.sdocx2 = StringVar()  # 设置所选路径进行显示
        self.edocx2 = tk.Entry(self.docxWindow, textvariable=self.sdocx2, width=30).grid(row=1, column=1, ipadx=5,ipady=2,pady=5)  # 设置entry为不可选不可写
        tk.Button(self.docxWindow, text='选择', command=lambda: self.openDirectory4(2)).grid(row=1, column=2, ipadx=5,ipady=2, pady=5, padx=5,sticky=tk.E)
        tk.Button(self.docxWindow, text='提交', command=lambda: self.docTodocxTask(mode)).grid(row=2, column=0,columnspan=3, ipadx=5,ipady=2, pady=5, padx=5)
        # 启动进度条,在提交按钮下显示进度条
        self.pb = ttk.Progressbar(self.docxWindow, length=200, value=0, mode="indeterminate")
        self.docxWindow.grab_set()

    # 进度条的开始方法
    def startProgressBar(self):
        self.pb.start(interval=5)

    # 选择一个要导入的路径doc转docx
    def openDirectory3(self, flag):
        path1 = askdirectory(title='选择一个路径')
        print path1
        if flag == 3:
            self.sdocx3.set(path1)
        if flag == 1:
            self.sdocx1.set(path1)

    # 选择一个要导出的路径doc转docx
    def openDirectory4(self, flag):
        path2 = askdirectory(title='选择一个路径')
        print path2
        if flag == 4:
            self.sdocx4.set(path2)
        if flag == 2:
            self.sdocx2.set(path2)

    # 弹出等待窗口,因为新窗口会循环所以必须使用新线程来创建
    def createWait(self):
        # print u"创建窗口方法执行"
        self.waitWindow = tk.Toplevel(self.master)
        # print u"创建了窗口对象"
        width = 225
        height = 200
        # 获取屏幕尺寸以计算布局参数，使窗口居屏幕中央
        screenwidth = self.waitWindow.winfo_screenwidth()
        screenheight = self.waitWindow.winfo_screenheight()
        alignstr = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
        self.waitWindow.title('处理中请稍候...')
        self.waitWindow.geometry(alignstr)
        self.waitWindow.resizable(0, 0)  # 防止用户调尺寸
        # 消除窗口的可选项(消除最小化，和退出)
        self.waitWindow.overrideredirect(True)
        self.waitWindow.attributes("-toolwindow", 1)
        self.waitWindow.wm_attributes("-topmost", 1)
        self.waitWindow.grab_set()
        # 扫描文件夹下的所有gif图片，然后得到图片对象放入到一个数组中
        ps = file_name("0")
        self.photos = []
        for pName in ps:
            self.photos.append(tk.PhotoImage(file="0/" + pName))
        canvas1 = tk.Canvas(self.waitWindow, width=225, height=200, bg="black")
        canvas1.pack()
        canvas1.create_text(112.5, 190, text='正在处理请稍候...', fill="white")
        # 调用定时回调函数执行图片切换
        self.update(0, canvas1, self.waitWindow)

    def shutWaitWindow(self):
        self.waitWindow.destroy()

    # 播放动画定时函数
    def update(self, idx, canvas, root):
        p = self.photos[idx]
        idx += 1
        # 一共三张图片,每次到第四次递归之前归零一下
        if idx == 3:
            idx = 0
        canvas.create_image(0, 0, anchor='nw', image=p)
        root.after(200, self.update, idx, canvas, root)

    # 检测处理完成定时函数
    def taskOver(self, root):
        if self.window_flag == 1:
            # print u"执行了销毁窗口"
            self.shutWaitWindow()
            # 等待动画结束之后跳出弹窗提示
            print self.window_flag
            self.complate_info()
            return  # 结束递归
        root.after(200, self.taskOver, root)

    # 检测doc转docx完毕定时函数
    def docTodocxOver(self, root):
        if self.doc_flag == 1:
            # 等待动画结束之后跳出弹窗提示
            print '转换结束标志位：' + str(self.doc_flag)
            # 消除窗口的不可选中状态
            # 隐藏进度条
            self.pb.stop()
            self.pb.grid_forget()
            self.docxWindow.attributes("-disabled", False)
            self.docxWindow.attributes("-topmost", False)
            self.complate_info()
            self.doc_flag = 0
            return  # 结束递归
        # 如果路径导入导出有一个无效则进行弹窗
        if self.doc_flag == 2:
            # 消除窗口的不可选中状态
            # 隐藏进度条
            self.pb.stop()
            self.pb.grid_forget()
            self.docxWindow.attributes("-disabled", False)
            self.docxWindow.attributes("-topmost", False)
            self.message_error()
            self.doc_flag = 0
            return  # 结束递归
        if self.doc_flag == 3:
            # 等待动画结束之后跳出弹窗提示
            print '转换结束标志位：' + str(self.doc_flag)
            # 消除窗口的不可选中状态
            # 隐藏进度条
            self.pb.stop()
            self.pb.grid_forget()
            self.WXPicWindow.attributes("-disabled", False)
            self.WXPicWindow.attributes("-topmost", False)
            self.complate_info()
            self.doc_flag = 0
            return  # 结束递归
        # 如果路径导入导出有一个无效则进行弹窗
        if self.doc_flag == 4:
            # 消除窗口的不可选中状态
            # 隐藏进度条
            self.pb.stop()
            self.pb.grid_forget()
            self.WXPicWindow.attributes("-disabled", False)
            self.WXPicWindow.attributes("-topmost", False)
            self.message_error()
            self.doc_flag = 0
            return  # 结束递归
        root.after(200, self.docTodocxOver, root)

    # 选择一个要导入的路径
    def openDirectory1(self):
        # 导入时判断一下模式
        if self.com.get() == 'xls/xlsx内容导入B':
            EB_fileName = askopenfilename(title='选择一个xls/xlsx', filetypes=[('XLS', '*.xls'), ('XLSX', '*.xlsx')])
            print EB_fileName
            self.importPath = EB_fileName
            self.pathEntry1_s.set(EB_fileName)  # 改变label的路径显示
        else:
            path1 = askdirectory(title='选择一个路径')
            print path1
            self.importPath = path1
            self.pathEntry1_s.set(path1)

    # 选择一个要导出的路径
    # 选择一个要导出的xls/xlsx的文件
    def openDirectory2(self):
        # 如果当前模式是Excel_A模式那么应该得到一个文件的路径且该文件必须是xls或者xlsx
        # 判断是否是Excel_A模式
        if self.com.get() == 'docx内容替换':
            path2 = askdirectory(title='选择一个路径')
            print path2
            self.exportPath = path2
            # self.pathL2['text'] = path2    #改变label的路径显示
            self.pathEntry2_s.set(path2)
        elif self.com.get() == 'xls/xlsx内容导入A':
            EA_fileName = askopenfilename(title='选择一个xls/xlsx', filetypes=[('XLS', '*.xls'), ('XLSX', '*.xlsx')])
            print EA_fileName
            self.exportPath = EA_fileName
            # self.pathL2['text'] = EA_fileName    #改变label的路径显示
            self.pathEntry2_s.set(EA_fileName)
        elif self.com.get() == 'xls/xlsx内容导入B':
            EB_fileName = askopenfilename(title='选择一个xls/xlsx', filetypes=[('XLS', '*.xls'), ('XLSX', '*.xlsx')])
            print EB_fileName
            self.exportPath = EB_fileName
            # self.pathL2['text'] = EB_fileName    #改变label的路径显示
            self.pathEntry2_s.set(EB_fileName)
            # path2 = askdirectory(title='选择一个路径')
            # print path2
            # self.exportPath = path2
            # self.pathL2['text'] = path2    #改变label的路径显示

    # 菜单栏点击保存事件处理函数
    def saveXml(self):
        # 打开对话框得到文件名名
        file1 = asksaveasfile(title='保存', initialfile='未命名.xml', filetypes=[('XML', '*.xml')]
                              , defaultextension='.xml')
        if file1 != None:
            print file1.name
            fileName1 = file1.name
            if fileName1 != None:
                # 创建文件之后使用多线程的方式来把数据写入xml
                # 根据模式来选择写入方式
                if self.com.get() == 'docx内容替换':
                    self.thread_window(self.writeDataXML, fileName1)
                elif self.com.get() == 'xls/xlsx内容导入A':
                    self.thread_window(self.writeDataXML_A, fileName1)
                elif self.com.get() == 'xls/xlsx内容导入B':
                    self.thread_window(self.writeDataXML_B, fileName1)

    # 菜单栏加载一个xml文件
    # 根据模式来加载对应的xml
    def loadXml(self):
        # 打开对话框得到文件名称
        fileName2 = askopenfilename(title='加载')
        if fileName2 != None and (os.path.exists(fileName2)):
            if self.com.get() == 'docx内容替换':
                print fileName2
                self.readDataXML(fileName2)
            elif self.com.get() == 'xls/xlsx内容导入A':
                print fileName2
                self.readDataXML_A(fileName2)
            elif self.com.get() == 'xls/xlsx内容导入B':
                print fileName2
                self.readDataXML_B(fileName2)

    # 工具栏的doc转docx功能
    def docTodocx(self):
        # 1弹出窗口2两个按钮选择路径3两个路径显示4最后一个提交按钮
        print '创建doc转docx窗口'
        self.createDocToDocx()  # 创建窗口

    # 帮助模式说明会自动打开电脑默认浏览器并进入我的简书
    def goToLink(self):
        webbrowser.open_new_tab('https://www.jianshu.com/u/a8beefd3c05f')

    # 创建一个新线程方法用来执行保存操作
    @staticmethod
    def thread_window(func, fileName):
        t = Thread(target=func, args=(fileName,))
        t.start()

    # 创建一个新线程方法用来执行处理任务
    @staticmethod
    def thread_task(func):
        t = Thread(target=func)
        t.start()

    # doc文档的xml保存操作
    def writeDataXML(self, fileName):
        # domTree = parse(u'./未命名.xml')
        # 1. 创建dom树对象
        domTree = minidom.Document()
        # 创建根节点并添加到dom树
        rootNode = domTree.createElement('content')
        domTree.appendChild(rootNode)
        # 创建entries节点
        entries_node = domTree.createElement('entries')
        # 创建entry操作，动态增加节点，遍历循环entry列表
        for num in range(0, len(self.entry_list) - 1, 2):
            # 输入框任一一栏没有输入内容不会执行替换操作
            if self.entry_list[num].get() != '' and self.entry_list[num + 1].get() != '':
                self.createNode(domTree, entries_node, self.entry_list[num].get().strip(),
                                self.entry_list[num + 1].get().strip())  # 编号年份
                print '保存old_text: ' + self.entry_list[num].get()
                print '保存new_next: ' + self.entry_list[num + 1].get()
        # 把entries节点放入父节点
        rootNode.appendChild(entries_node)
        # 把创建号的节点写入到文件
        with open(fileName, 'wb') as f:
            # 缩进 - 换行 - 编码
            # indent='', addindent='\t',newl ='\n'
            domTree.writexml(f, addindent='\t', newl='\n', encoding="utf-8")
            # 这里添加换行的话会在原先的xml里给所有元素添加换行，而且每次都会添加

    # excel模式B的保存操作
    def writeDataXML_B(self, fileName):
        # domTree = parse(u'./未命名.xml')
        # 1. 创建dom树对象
        domTree = minidom.Document()
        # 创建根节点并添加到dom树
        rootNode = domTree.createElement('content')
        domTree.appendChild(rootNode)
        ##获取开始复制的行和列写入xml
        self.createNode_B(domTree, rootNode, self.strB_2.get(), self.strB_3.get())
        # 创建entries节点
        entries_node = domTree.createElement('entries')
        # 创建entry操作，动态增加节点，遍历循环entry列表
        for num in range(0, len(self.entry_list_B) - 3, 4):
            # 只读入有效坐标只有行或列或者全没有不读入直接跳入下一个循环
            if (self.entry_list_B[num + 1].get() != '' and self.entry_list_B[num + 2].get() != '') and (
                    self.entry_list_B[num + 1].get() != None and self.entry_list_B[num + 2].get() != None):
                t = self.entry_list_B[num].get().strip()  # 第一个是坐标提示
                cc = self.entry_list_B[num + 1].get().strip()  # 第二个抓取列
                r = self.entry_list_B[num + 2].get().strip()  # 第三个复制行
                c = self.entry_list_B[num + 3].get().strip()  # 第四个复制列
                print '保存r: ' + r
                print '保存c: ' + c
                self.createNode_position_B(domTree, entries_node, cc, r, c, t)
            else:
                continue
        # 把entries节点放入父节点
        rootNode.appendChild(entries_node)
        # 把创建号的节点写入到文件
        with open(fileName, 'wb') as f:
            # 缩进 - 换行 - 编码
            # indent='', addindent='\t',newl ='\n'
            domTree.writexml(f, addindent='\t', newl='\n', encoding="utf-8")
            # 这里添加换行的话会在原先的xml里给所有元素添加换行，而且每次都会添加

    # excel模式A的保存操作
    def writeDataXML_A(self, fileName):
        # domTree = parse(u'./未命名.xml')
        # 1. 创建dom树对象
        domTree = minidom.Document()
        # 创建根节点并添加到dom树
        rootNode = domTree.createElement('content')
        domTree.appendChild(rootNode)
        ##获取开始复制的行和列写入xml
        self.createNode_A(domTree, rootNode, self.s2_1.get(), self.s2_2.get())
        # 创建entries节点
        entries_node = domTree.createElement('entries')
        # 创建entry操作，动态增加节点，遍历循环entry列表
        for num in range(0, len(self.entry_list_A) - 2, 3):
            # 只读入有效坐标只有行或列或者全没有不读入直接跳入下一个循环
            if (self.entry_list_A[num + 1].get() != '' and self.entry_list_A[num + 2].get() != '') and (
                    self.entry_list_A[num + 1].get() != None and self.entry_list_A[num + 2].get() != None):
                t = self.entry_list_A[num].get().strip()  # 第一个是坐标提示
                r = self.entry_list_A[num + 1].get().strip()  # 第二个是行
                c = self.entry_list_A[num + 2].get().strip()  # 第三个是列
                print '保存r: ' + r
                print '保存c: ' + c
                self.createNode_position(domTree, entries_node, r, c, t)
            else:
                continue
        # 把entries节点放入父节点
        rootNode.appendChild(entries_node)
        # 把创建号的节点写入到文件
        with open(fileName, 'wb') as f:
            # 缩进 - 换行 - 编码
            # indent='', addindent='\t',newl ='\n'
            domTree.writexml(f, addindent='\t', newl='\n', encoding="utf-8")
            # 这里添加换行的话会在原先的xml里给所有元素添加换行，而且每次都会添加

    # 读取xml文件下的数据
    def readDataXML(self, fileName):
        # 重置组件
        self.entry_row = 0
        self.c_height = 150
        domTree = parse(fileName)
        # 文档根元素
        rootNode = domTree.documentElement
        # print rootNode.nodeName
        # 所有输入框内容
        entries = rootNode.getElementsByTagName("entry")
        self.scroll_len_doc = 170
        # 判断entryies是否为空,然后开始销毁组件
        if entries != None:
            for (label, entry) in zip(self.entry_list, self.entry_label):
                label.destroy()
                entry.destroy()
            # 重置组件的列表，全部置为空列表
            self.entry_list = []
            self.entry_label = []
            self.entry_str = []
            for entry in entries:
                if entry.hasAttribute("ID"):
                    # print("ID:", entry.getAttribute("ID"))  #获取属性
                    if entry.firstChild != None:  # 获取数据
                        # print entry.childNodes.count
                        if self.scroll_len_doc == 170:
                            # 写一个switch把获取到的值放入entry里面就行了
                            self.switch_entry(entry.getAttribute("ID"), entry.childNodes[0].data, 20)
                        else:
                            self.switch_entry(entry.getAttribute("ID"), entry.childNodes[0].data, 30)
                    else:
                        self.switch_entry(entry.getAttribute("ID"), '')

    # 进行载入的字符载入操作和生成控件操作
    def switch_entry(self, id, data, s_length):
        # 填入id和data每填入一次生成一次组件
        # 调用添加组件方法
        self.b_add(s_length)
        # 给生成的entry进行赋值操作，取除列表的后两个
        self.entry_str[len(self.entry_str) - 1].set(data)
        self.entry_str[len(self.entry_str) - 2].set(id)

    # 读取excel模式A的xml文件下的数据
    def readDataXML_B(self, fileName):
        # 重置组件
        self.EB_entry_row = 2  # 加载是会把列表清空所以从第一个开始
        self.EB_height = 140 + 30  # 在载入时把滑块给归零
        # 设置滚动的滚动范围，发现范围显示不全((0,0,self.c_height,self.c_height))或者canvas.bbox('all')
        # self.canvas1.configure(scrollregion = ((0,0,self.EA_height,self.EA_height)))#将滚动区域设置为画布边界
        domTree = parse(fileName)
        # 文档根元素
        rootNode = domTree.documentElement
        # print rootNode.nodeName
        # 获得开始的行数和列数,放入对应输入框
        crawlRow_node = rootNode.getElementsByTagName("CrawlRow")
        fileReName_node = rootNode.getElementsByTagName("FileReName")
        self.strB_2.set(crawlRow_node[0].firstChild.data)
        self.strB_3.set(fileReName_node[0].firstChild.data)
        # 获得所有输入坐标和提示内容
        entries = rootNode.getElementsByTagName("entry")
        # 判断entryies是否为空,然后开始销毁组件
        if entries != None:
            for (label, entry) in zip(self.entry_list_B, self.entry_label_B):
                label.destroy()
                entry.destroy()
            # 重置组件的列表，全部置为空列表
            self.entry_list_B = []
            self.entry_label_B = []
            self.entry_str_B = []
            self.scroll_len = 170
            for entry in entries:
                if entry.firstChild != None:  # 获取数据
                    cc = entry.getElementsByTagName("crawlC")[0].firstChild.data
                    # print '读取数据row --->' + row
                    row = entry.getElementsByTagName("row")[0].firstChild.data
                    col = entry.getElementsByTagName("col")[0].firstChild.data
                    # 如果用户没有输入对应提示则默认是空字符串
                    tip = entry.getElementsByTagName("tip")[0].firstChild.data
                    if tip == None:
                        tip = ''
                    if self.scroll_len == 170:
                        self.switch_entry_B(tip, cc, row, col, 20)
                    else:
                        self.switch_entry_B(tip, cc, row, col, 30)
                else:
                    self.switch_entry_B('', '', '', '')
            # 每次遍历完xml都要把s_length置为0
            self.scroll_len = 0

    # 进行载入的字符载入操作和生成控件操作
    def switch_entry_B(self, e1, e2, e3, e4, s_length):
        # 填入id和data每填入一次生成一次组件
        # 调用添加组件方法
        # 第一次添加20其余都是30
        self.b_add(s_length)
        # 给生成的entry进行赋值操作，取列表的后两个
        self.entry_str_B[len(self.entry_str_B) - 4].set(e1)
        self.entry_str_B[len(self.entry_str_B) - 3].set(e2)
        self.entry_str_B[len(self.entry_str_B) - 2].set(e3)
        self.entry_str_B[len(self.entry_str_B) - 1].set(e4)

    # 读取excel模式A的xml文件下的数据
    def readDataXML_A(self, fileName):
        # 重置组件
        self.EA_entry_row = 2  # 加载是会把列表清空所以从第一个开始
        self.EA_height = 140 + 30  # 在载入时把滑块给归零
        # 设置滚动的滚动范围，发现范围显示不全((0,0,self.c_height,self.c_height))或者canvas.bbox('all')
        # self.canvas1.configure(scrollregion = ((0,0,self.EA_height,self.EA_height)))#将滚动区域设置为画布边界
        domTree = parse(fileName)
        # 文档根元素
        rootNode = domTree.documentElement
        # print rootNode.nodeName
        # 获得开始的行数和列数,放入对应输入框
        startRow_node = rootNode.getElementsByTagName("startRow")
        startCol_node = rootNode.getElementsByTagName("startCol")
        self.s2_1.set(startRow_node[0].firstChild.data)
        self.s2_2.set(startCol_node[0].firstChild.data)
        # 获得所有输入坐标和提示内容
        entries = rootNode.getElementsByTagName("entry")
        # 判断entryies是否为空,然后开始销毁组件
        if entries != None:
            for (label, entry) in zip(self.entry_list_A, self.entry_label_A):
                label.destroy()
                entry.destroy()
            # 重置组件的列表，全部置为空列表
            self.entry_list_A = []
            self.entry_label_A = []
            self.entry_str_A = []
            self.scroll_len = 170
            for entry in entries:
                if entry.firstChild != None:  # 获取数据
                    row = entry.getElementsByTagName("row")[0].firstChild.data
                    # print '读取数据库row --->' + row
                    col = entry.getElementsByTagName("col")[0].firstChild.data
                    # 如果用户没有输入对应提示则默认是空字符串
                    tip = entry.getElementsByTagName("tip")[0].firstChild.data
                    if tip == None:
                        tip = ''
                    if self.scroll_len == 170:
                        self.switch_entry_A(tip, row, col, 20)
                    else:
                        self.switch_entry_A(tip, row, col, 30)
                else:
                    self.switch_entry_A('', '', '')
            # 每次遍历完xml都要把s_length置为0
            self.scroll_len = 0

    # 进行载入的字符载入操作和生成控件操作
    def switch_entry_A(self, e1, e2, e3, s_length):
        # 填入id和data每填入一次生成一次组件
        # 调用添加组件方法
        # 第一次添加20其余都是30
        self.b_add(s_length)
        # 给生成的entry进行赋值操作，取列表的后两个
        self.entry_str_A[len(self.entry_str_A) - 3].set(e1)
        self.entry_str_A[len(self.entry_str_A) - 2].set(e2)
        self.entry_str_A[len(self.entry_str_A) - 1].set(e3)

    # dom树对象,父节点对象，序号，组件值
    def createNode(self, domTree, entries_node, num, kitData):
        # 新建一个entry节点
        entry_node = domTree.createElement("entry")
        # 设置属性
        entry_node.setAttribute("ID", num)
        # 创建文本节点
        entry_text_value = domTree.createTextNode(kitData)  # 这里不需要添加u
        # 把文本节点挂载到entry节点
        entry_node.appendChild(entry_text_value)
        # 把子节点放入到父节点
        entries_node.appendChild(entry_node)

    # 模式A创建复制行和列节点创建
    def createNode_B(self, domTree, parent_node, rowData, colData):
        # 新建一个entry节点
        startRow_node = domTree.createElement('CrawlRow')
        startCol_node = domTree.createElement('FileReName')
        # 创建文本节点
        startRow_node_text_value = domTree.createTextNode(rowData)  # 这里不需要添加u
        startCol_node_text_value = domTree.createTextNode(colData)  # 这里不需要添加u
        # 把文本节点挂载到节点
        startRow_node.appendChild(startRow_node_text_value)
        startCol_node.appendChild(startCol_node_text_value)
        # 把子节点放入到父节点
        parent_node.appendChild(startRow_node)
        parent_node.appendChild(startCol_node)

    # 模式A创建复制行和列节点创建
    def createNode_A(self, domTree, parent_node, rowData, colData):
        # 新建一个entry节点
        startRow_node = domTree.createElement('startRow')
        startCol_node = domTree.createElement('startCol')
        # 创建文本节点
        startRow_node_text_value = domTree.createTextNode(rowData)  # 这里不需要添加u
        startCol_node_text_value = domTree.createTextNode(colData)  # 这里不需要添加u
        # 把文本节点挂载到节点
        startRow_node.appendChild(startRow_node_text_value)
        startCol_node.appendChild(startCol_node_text_value)
        # 把子节点放入到父节点
        parent_node.appendChild(startRow_node)
        parent_node.appendChild(startCol_node)

    # 模式A创建坐标存储单元
    def createNode_position(self, domTree, parent_node, rowData, colData, tip):
        # 新建一个entry节点
        entry_node = domTree.createElement("entry")
        # 一个entry包含三个子节点
        entry_node_r = domTree.createElement("row")
        entry_node_c = domTree.createElement("col")
        entry_node_t = domTree.createElement("tip")
        # 创建文本节点
        entry_node_r_t = domTree.createTextNode(rowData)  # 这里不需要添加u
        entry_node_c_t = domTree.createTextNode(colData)
        entry_node_t_t = domTree.createTextNode(tip)
        # 把文本节点挂载到entry节点
        entry_node_r.appendChild(entry_node_r_t)
        entry_node_c.appendChild(entry_node_c_t)
        entry_node_t.appendChild(entry_node_t_t)
        entry_node.appendChild(entry_node_r)
        entry_node.appendChild(entry_node_c)
        entry_node.appendChild(entry_node_t)
        # 把子节点放入到父节点
        parent_node.appendChild(entry_node)

    # 模式B创建坐标存储单元
    def createNode_position_B(self, domTree, parent_node, crawlCol, rowData, colData, tip):
        # 新建一个entry节点
        entry_node = domTree.createElement("entry")
        # 一个entry包含三个子节点
        entry_node_cc = domTree.createElement("crawlC")
        entry_node_r = domTree.createElement("row")
        entry_node_c = domTree.createElement("col")
        entry_node_t = domTree.createElement("tip")
        # 创建文本节点
        entry_node_cc_t = domTree.createTextNode(crawlCol)
        entry_node_r_t = domTree.createTextNode(rowData)  # 这里不需要添加u
        entry_node_c_t = domTree.createTextNode(colData)
        entry_node_t_t = domTree.createTextNode(tip)
        # 把文本节点挂载到entry节点
        entry_node_cc.appendChild(entry_node_cc_t)
        entry_node_r.appendChild(entry_node_r_t)
        entry_node_c.appendChild(entry_node_c_t)
        entry_node_t.appendChild(entry_node_t_t)
        entry_node.appendChild(entry_node_cc)
        entry_node.appendChild(entry_node_r)
        entry_node.appendChild(entry_node_c)
        entry_node.appendChild(entry_node_t)
        # 把子节点放入到父节点
        parent_node.appendChild(entry_node)


if __name__ == '__main__':
    window = tk.Tk()
    width = 650
    height = 350
    # 获取屏幕尺寸以计算布局参数，使窗口居屏幕中央
    screenwidth = window.winfo_screenwidth()
    screenheight = window.winfo_screenheight()
    alignstr = '%dx%d+%d+%d' % (width, height, (screenwidth - width) / 2, (screenheight - height) / 2)
    window.title('SeekLife自动文档工具v1.0')  # 设置窗口标题
    window.geometry(alignstr)  # 设置窗口在屏幕的位置
    window.resizable(0, 0)  # 防止用户调尺寸
    window.iconbitmap('icons/aaa12.ico')
    app = Application(master=window)
    app.mainloop()
